# -*- coding: utf-8 -*-
"""
md_to_docx.py - 보고서 마크다운(.md)을 워드(.docx)로 변환

실행:
    py -3 md_to_docx.py <입력.md> <출력.docx>

지원 문법(이 스킬의 보고서 형식에 맞춤):
  #/##/### 제목, - 또는 * 불릿, 1. 번호목록, | 표 |, **굵게**, *기울임*, --- 구분선
python-docx 필요(없으면 자동 설치 시도). Word가 없어도 .docx 생성 가능(순수 파이썬).
"""
import sys
import re
import subprocess


def ensure_docx():
    try:
        import docx  # noqa
    except ImportError:
        subprocess.run([sys.executable, "-m", "pip", "install", "--quiet", "python-docx"], check=True)


def add_runs(paragraph, text):
    """**굵게**, *기울임* 인라인 처리."""
    # **bold** 를 먼저 토큰화, 그 안팎에서 *italic* 처리
    for chunk in re.split(r"(\*\*[^*]+\*\*)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            paragraph.add_run(chunk[2:-2]).bold = True
        else:
            for sub in re.split(r"(\*[^*]+\*)", chunk):
                if not sub:
                    continue
                if sub.startswith("*") and sub.endswith("*") and len(sub) > 2:
                    paragraph.add_run(sub[1:-1]).italic = True
                else:
                    paragraph.add_run(sub)


def is_table_row(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def parse_cells(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep_row(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def main():
    if len(sys.argv) < 3:
        print("사용법: py -3 md_to_docx.py <입력.md> <출력.docx>", file=sys.stderr)
        sys.exit(2)
    ensure_docx()
    from docx import Document
    from docx.shared import Pt

    src, out = sys.argv[1], sys.argv[2]
    with open(src, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # 표 블록
        if is_table_row(line):
            block = []
            while i < len(lines) and is_table_row(lines[i]):
                if not is_sep_row(lines[i]):
                    block.append(parse_cells(lines[i]))
                i += 1
            if block:
                cols = max(len(r) for r in block)
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(block):
                    cells = table.add_row().cells
                    for c in range(cols):
                        txt = row[c] if c < len(row) else ""
                        cells[c].text = ""
                        p = cells[c].paragraphs[0]
                        add_runs(p, txt)
                        if r == 0:
                            for run in p.runs:
                                run.bold = True
            continue

        if not s:
            i += 1
            continue
        if re.match(r"^-{3,}$", s):  # 구분선
            i += 1
            continue

        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif re.match(r"^\d+\.\s+", s):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", s))
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, s[2:])
        else:
            p = doc.add_paragraph()
            add_runs(p, s)
        i += 1

    doc.save(out)
    print(f"워드 문서 저장 완료: {out}")


if __name__ == "__main__":
    main()
